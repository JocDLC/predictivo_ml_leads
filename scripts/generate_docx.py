"""
Script para generar el documento técnico 'Documentación Técnica y Manual del Usuario.docx'
con toda la información consolidada del proyecto Predictivo ML Leads V2.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_PATH = os.path.join(BASE_DIR, "Doc_Tecnica_v2.docx")


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a formatted table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "2B579A")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, value in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(value)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    return table


def add_kv_table(doc, pairs):
    """Two-column key-value table."""
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (key, val) in enumerate(pairs):
        ck = table.rows[i].cells[0]
        cv = table.rows[i].cells[1]
        ck.text = key
        cv.text = str(val)
        for p in ck.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        for p in cv.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
        set_cell_shading(ck, "EBF5FB")
    return table


def build_document():
    doc = Document()

    # --- Page style ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.color.rgb = RGBColor(43, 87, 154)

    # ================================================================
    # PORTADA
    # ================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\nDocumentación Técnica y Manual del Usuario")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(43, 87, 154)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Predictivo ML Leads\nClasificación de Hot/Cold Leads para Renault México")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(80, 80, 80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nFuente única de verdad: estado vigente del repositorio y artefactos locales.\nMayo 2026 — Versión 2.2")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # --- Ficha técnica ---
    doc.add_heading("Ficha del proyecto", level=2)
    add_kv_table(doc, [
        ("Nombre del proyecto", "Predictivo ML Leads — Clasificación de Hot/Cold Leads para Renault México"),
        ("Tipo de modelo", "Clasificación binaria supervisada"),
        ("Modelo desplegado", "GradientBoostingClassifier (scikit-learn)"),
        ("Umbral operativo", "0.325"),
        ("Fecha de actualización", "Mayo 2026"),
        ("Versión del documento", "v2.2"),
        ("Dataset analítico V2", "58,546 leads limpios de 2025 (excluyendo Abr/May/Jun, chatbot, Arkana/Oroch)"),
        ("Variable objetivo", "target (1 = Hot Lead / Contacto interesado, 0 = Cold Lead)"),
        ("Demo / despliegue", "App local Streamlit. Ejecutar: .\\.venv\\Scripts\\python.exe -m streamlit run streamlit_app/app.py"),
        ("Repositorio", "https://github.com/JocDLC/predictivo_ml_leads"),
    ])

    doc.add_page_break()

    # ================================================================
    # 1. RESUMEN EJECUTIVO
    # ================================================================
    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "El sistema clasifica automáticamente cada lead nuevo en Hot o Cold y devuelve una probabilidad de conversión "
        "para priorizar la atención comercial. La versión desplegada en la app usa un GradientBoostingClassifier con "
        "22 features finales y un umbral operativo de 0.325."
    )
    doc.add_paragraph(
        "El objetivo del umbral 0.325 es capturar más Hot Leads reales sin disparar de forma descontrolada los falsos "
        "positivos. Frente al umbral por defecto de 0.50, este ajuste recupera 1,191 Hot Leads adicionales en test "
        "(FN baja de 2,457 a 1,266), a costa de 1,924 Cold Leads extras priorizados como falsos positivos."
    )
    doc.add_paragraph(
        "Métricas vigentes en test a umbral 0.325: ROC-AUC 0.7828, Accuracy 0.6947, Precision 0.5604, "
        "Recall 0.6992 y F1 0.6221."
    )

    # ================================================================
    # 2. PROBLEMA Y CONTEXTO
    # ================================================================
    doc.add_heading("2. Problema y contexto", level=1)
    doc.add_paragraph(
        "Renault México recibe leads desde formularios web, plataformas de marketing y otras fuentes digitales. "
        "Sin un mecanismo automático de priorización, el equipo comercial dedica tiempo equivalente a oportunidades "
        "con alta intención de compra y a contactos fríos o de baja calidad."
    )
    doc.add_paragraph(
        "El predictivo busca reducir ese costo operativo: prioriza leads con mayor probabilidad de conversión, "
        "mejora la velocidad de respuesta y documenta por qué una observación fue clasificada como Hot o Cold "
        "mediante explicaciones SHAP."
    )

    # ================================================================
    # 3. ARQUITECTURA DEL SISTEMA
    # ================================================================
    doc.add_heading("3. Arquitectura del sistema", level=1)

    doc.add_heading("3.1 Flujo técnico de inferencia", level=2)
    doc.add_paragraph(
        "Archivo Excel → streamlit_app/app.py → streamlit_app/core/inference.py → detección de columnas por keywords → "
        "extracción de fecha y normalización → limpieza de nulos → feature engineering alineado al artefacto → "
        "predicción predict_proba → aplicación del umbral 0.325 → resultado Hot/Cold con probabilidad."
    )
    doc.add_paragraph(
        "La inferencia no depende de un esquema rígido de 28 columnas exactas: el parser detecta columnas "
        "semánticas por nombre, deriva mes/día/hora/día_semana desde la fecha y luego reordena el DataFrame "
        "para coincidir con las 22 features esperadas por el modelo."
    )

    doc.add_heading("3.2 Persistencia", level=2)
    doc.add_paragraph(
        "La app usa SQLite en data/memory/history.db mediante streamlit_app/core/memory.py. "
        "Allí guarda sesiones, metadatos y predicciones individuales para la página Historial."
    )

    doc.add_heading("3.3 Páginas de la app", level=2)
    bullets = [
        "Predicción de Leads: carga de Excel, inferencia, grilla de resultados con filtros, detalle individual con SHAP.",
        "Reporte del Modelo: 5 tabs (Data Engineering, EDA, Feature Engineering, Modelado, Evaluación) con gráficas interactivas Plotly y explicaciones de negocio.",
        "Historial: sesiones previas guardadas en SQLite, tendencia de % Hot, exportación CSV.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # ================================================================
    # 4. DATOS Y VARIABLES
    # ================================================================
    doc.add_heading("4. Datos y variables", level=1)

    doc.add_heading("4.1 Pipeline de datos V2", level=2)
    doc.add_paragraph("Versión de verdad adoptada en este documento:")
    data_bullets = [
        "Fuente cruda: Excel Salesforce CRM con 906,894 registros útiles y 28 columnas.",
        "Corte analítico V2: solo año 2025, excluyendo abril, mayo y junio por anomalías operativas.",
        "Se eliminan 149,704 leads del chatbot (plataforma CHATBOT) para modelar solo gestión humana.",
        "Se eliminan 7,566 leads sin cualificación (no permiten construir el target).",
        "Se eliminan 4,123 leads preclasificados como Hot de los formularios Arkana y Oroch (data leakage validado con equipo comercial, abril 2026).",
        "Se eliminan 3,359 duplicados.",
        "Dataset limpio final: 58,546 filas × 12 columnas.",
        "Distribución final: 21,042 Hot (35.94%) y 37,504 Cold (64.06%).",
        "Split de modelado: train 46,836 filas y test 11,710 filas (80/20 estratificado, random_state=42).",
    ]
    for b in data_bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("4.2 Funnel de limpieza documentado (notebook 00)", level=2)
    add_styled_table(doc,
        ["Etapa", "Registros"],
        [
            ["Leads 2025", "322,009"],
            ["Excluir Abr/May/Jun 2025", "223,298"],
            ["Eliminar chatbot", "73,594"],
            ["Eliminar sin cualificación", "66,028"],
            ["Eliminar leakage Arkana/Oroch", "61,905"],
            ["Eliminar duplicados → dataset final", "58,546"],
        ],
    )

    doc.add_heading("4.3 Variables incluidas en el pipeline vigente", level=2)
    add_styled_table(doc,
        ["Variable", "Estado", "Uso actual"],
        [
            ["nombre_formulario", "Incluida", "Feature final del modelo."],
            ["campana", "Incluida", "Feature final del modelo."],
            ["origen_creacion", "Incluida", "One-Hot Encoded (IGNITIONONE, LOCAL MX, ONE, RRSS)."],
            ["vehiculo_interes", "Incluida", "Feature de negocio; entre las más influyentes."],
            ["concesionario", "Incluida", "Feature vigente; Target Encoded."],
            ["origen", "Incluida", "Feature vigente; top global de importancia."],
            ["mes_creacion", "Incluida", "Variable temporal; feature más importante del modelo."],
            ["dia_creacion", "Incluida", "Variable temporal derivada."],
            ["hora_creacion", "Incluida", "Variable temporal derivada."],
            ["dia_semana_creacion", "Incluida", "One-Hot Encoded (lunes a domingo)."],
            ["es_fin_de_semana", "Derivada", "Binaria: 1 si sábado/domingo."],
            ["franja_horaria", "Derivada", "Madrugada/mañana/tarde/noche; One-Hot Encoded."],
        ],
    )

    doc.add_heading("4.4 Variables excluidas o controladas", level=2)
    add_styled_table(doc,
        ["Variable / Regla", "Estado", "Motivo"],
        [
            ["plataforma / MX_LEAD_QUALIF", "Excluida", "Data leakage severo; la cola CRM reflejaba clasificación posterior."],
            ["Arkana y Oroch (formularios)", "Excluidos", "Leads enviados como Hot sin cualificación humana; leakage validado."],
            ["Leads de chatbot", "Excluidos", "No representan la operación humana que se quiere priorizar."],
            ["Abril, mayo y junio 2025", "Excluidos", "Meses con comportamiento operativo anómalo."],
            ["subtipo_interes", "Excluida", "Riesgo de data leakage; refleja información post-clasificación."],
            ["lead_id", "Excluida del modelado", "Trazabilidad, sin valor predictivo directo."],
            ["PII (nombre, teléfono, correo)", "Excluidas", "Privacidad y falta de robustez para inferencia estructurada."],
            ["nombre_formulario vs campaña", "Pendiente abierto", "Redundancia 1:1 para formularios MX_Renault; el modelo actual usa ambas."],
        ],
    )

    doc.add_heading("4.5 Features finales del modelo (22)", level=2)
    doc.add_paragraph(
        "El modelo desplegado espera exactamente estas 22 features finales (en el orden de feature_names_in_): "
        "nombre_formulario, campana, vehiculo_interes, concesionario, origen, mes_creacion, dia_creacion, "
        "hora_creacion, es_fin_de_semana, origen_creacion_IGNITIONONE, origen_creacion_LOCAL MX, origen_creacion_ONE, "
        "origen_creacion_RRSS, dia_semana_creacion_jueves, dia_semana_creacion_lunes, dia_semana_creacion_martes, "
        "dia_semana_creacion_miercoles, dia_semana_creacion_sabado, dia_semana_creacion_viernes, "
        "franja_horaria_manana, franja_horaria_noche, franja_horaria_tarde."
    )

    # ================================================================
    # 5. FEATURE ENGINEERING
    # ================================================================
    doc.add_heading("5. Feature Engineering", level=1)

    doc.add_heading("5.1 Transformaciones aplicadas", level=2)
    fe_steps = [
        "Derivación temporal: a partir de la fecha de creación se extraen mes_creacion, dia_creacion, hora_creacion, dia_semana_creacion.",
        "Feature es_fin_de_semana: binaria (1 si sábado/domingo, 0 en caso contrario).",
        "Feature franja_horaria: clasifica la hora en madrugada (0-5), mañana (6-11), tarde (12-17) y noche (18-23).",
        "Eliminación de subtipo_interes por riesgo de data leakage.",
        "Agrupación de categorías: las categorías con frecuencia < 1% se agrupan como 'otros' para evitar overfitting.",
        "Target Encoding con smoothing bayesiano: para variables de alta cardinalidad (vehiculo_interes, origen, nombre_formulario, campana, concesionario). "
        "Reemplaza cada categoría por una tasa de conversión suavizada entre la media local y la media global.",
        "One-Hot Encoding: para variables de baja cardinalidad (origen_creacion, dia_semana_creacion, franja_horaria).",
        "Alineación final: el DataFrame se reindexa a las 22 features exactas esperadas por el modelo.",
    ]
    for s in fe_steps:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("5.2 ¿Por qué Bayesian Target Encoding?", level=2)
    doc.add_paragraph(
        "Variables como nombre_formulario o campaña tienen cientos de categorías únicas. Un One-Hot puro "
        "generaría 200+ columnas, muchas con muy pocos ejemplos, causando explosión de dimensionalidad y "
        "sobreajuste. El Target Encoding con smoothing bayesiano reduce cada variable a una sola columna "
        "numérica: una tasa de conversión suavizada entre la señal local (media de la categoría) y la señal "
        "estable (media global del dataset). Esto controla el ruido en categorías raras."
    )

    doc.add_heading("5.3 Split Train/Test", level=2)
    doc.add_paragraph(
        "Split estratificado 80/20 manteniendo la distribución del target (~35.94% Hot) en ambos conjuntos. "
        "Train: 46,836 filas. Test: 11,710 filas. El encoding se ajusta SOLO sobre train para evitar data leakage."
    )

    # ================================================================
    # 6. MODELO Y ENTRENAMIENTO
    # ================================================================
    doc.add_heading("6. Modelo y entrenamiento", level=1)

    doc.add_paragraph(
        "Narrativa adoptada: se documenta el modelo realmente desplegado. El artefacto models/best_model.joblib "
        "carga un GradientBoostingClassifier. Se elimina la narrativa anterior de Random Forest."
    )

    doc.add_heading("6.1 Comparación de modelos (notebook 03)", level=2)
    doc.add_paragraph(
        "Se evaluaron 4 algoritmos con las mismas 22 features y el mismo split train/test:"
    )
    add_styled_table(doc,
        ["Modelo", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC", ""],
        [
            ["Logistic Regression", "0.704", "0.633", "0.431", "0.513", "0.742", ""],
            ["Random Forest", "0.733", "0.667", "0.503", "0.573", "0.770", ""],
            ["Gradient Boosting", "0.738", "0.676", "0.517", "0.586", "0.779", "★ Ganador"],
            ["LightGBM", "0.736", "0.671", "0.512", "0.581", "0.777", ""],
        ],
    )
    doc.add_paragraph(
        "Gradient Boosting fue seleccionado por el mejor balance de métricas, especialmente ROC-AUC y F1."
    )

    doc.add_heading("6.2 Hiperparámetros del modelo desplegado", level=2)
    add_kv_table(doc, [
        ("n_estimators", "200"),
        ("learning_rate", "0.1"),
        ("max_depth", "5"),
        ("min_samples_leaf", "10"),
        ("loss", "log_loss"),
        ("criterion", "friedman_mse"),
        ("subsample", "1.0"),
        ("random_state", "42"),
    ])

    doc.add_heading("6.3 Validación cruzada", level=2)
    doc.add_paragraph(
        "Validación cruzada 5-fold sobre train: ROC-AUC = 0.7791 ± 0.0049. "
        "La desviación pequeña indica comportamiento estable entre particiones."
    )

    doc.add_heading("6.4 Procedimiento reproducible", level=2)
    steps = [
        "Limpieza y recorte operativo en notebooks/00_data_engineering_v2_2025.ipynb.",
        "Análisis exploratorio en notebooks/01_exploratory_data_analysis_v2.ipynb.",
        "Feature engineering y split estratificado 80/20 en notebooks/02_feature_engineering_v2.ipynb.",
        "Entrenamiento del mejor modelo en notebooks/03_modeling_v2.ipynb.",
        "Evaluación y SHAP contra best_model.joblib en notebooks/04_evaluation_v2.ipynb.",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    doc.add_paragraph(
        "Nota operativa: preprocessing_config.joblib aún conserva umbral legacy 0.35, pero la app resuelve "
        "ese caso y fuerza 0.325 como umbral activo en streamlit_app/core/inference.py."
    )

    # ================================================================
    # 7. MÉTRICAS Y RESULTADOS
    # ================================================================
    doc.add_heading("7. Métricas y resultados", level=1)
    doc.add_paragraph(
        "Métricas reproducibles sobre data/processed/X_test_v2.csv y y_test_v2.csv, con umbral 0.325."
    )

    add_styled_table(doc,
        ["Métrica", "Valor"],
        [
            ["ROC-AUC", "0.7828"],
            ["Accuracy", "0.6947"],
            ["Precision (Hot)", "0.5604"],
            ["Recall (Hot)", "0.6992"],
            ["F1 (Hot)", "0.6221"],
            ["Leads Hot reales en test", "4,209"],
            ["Leads Cold reales en test", "7,501"],
            ["Leads predichos Hot", "5,252"],
            ["Leads predichos Cold", "6,458"],
        ],
    )

    doc.add_heading("7.1 Matriz de confusión a umbral 0.325", level=2)
    add_styled_table(doc,
        ["", "Pred Cold", "Pred Hot"],
        [
            ["Real Cold", "5,192 (TN)", "2,309 (FP)"],
            ["Real Hot", "1,266 (FN)", "2,943 (TP)"],
        ],
    )

    doc.add_heading("7.2 Comparación de umbrales relevantes", level=2)
    add_styled_table(doc,
        ["Umbral", "Precision", "Recall", "F1", "Pred Hot", "Estrategia"],
        [
            ["0.300", "0.5308", "0.7510", "0.6220", "5,955", "Alternativa baja"],
            ["0.325", "0.5604", "0.6992", "0.6221", "5,252", "Activo ★"],
            ["0.350", "0.5940", "0.6500", "0.6208", "4,606", "Alternativa alta"],
            ["0.500", "0.8198", "0.4163", "0.5522", "2,137", "Baseline"],
        ],
    )
    doc.add_paragraph(
        "Lectura de negocio del umbral 0.325: es el compromiso operativo elegido entre 0.30 y 0.35. "
        "Contra 0.35 recupera más Hot Leads sin deteriorar F1; contra 0.50 sacrifica precisión pero "
        "evita perder una gran cantidad de oportunidades calientes reales. El F1-Score alcanza su máximo "
        "global precisamente en 0.325."
    )

    # ================================================================
    # 8. INTERPRETABILIDAD (SHAP)
    # ================================================================
    doc.add_heading("8. Interpretabilidad (SHAP)", level=1)
    doc.add_paragraph(
        "La explicabilidad vigente debe leerse sobre el GradientBoostingClassifier actual. "
        "SHAP se calcula en streamlit_app/core/shap_explainer.py para cada lead individual y en "
        "streamlit_app/components/model_report.py para análisis global."
    )

    doc.add_heading("8.1 Top features globales", level=2)
    doc.add_paragraph(
        "Top features por impacto SHAP promedio: mes_creacion, vehiculo_interes, origen, dia_creacion, "
        "es_fin_de_semana, concesionario, hora_creacion, nombre_formulario y campana."
    )
    doc.add_paragraph(
        "Las importancias Gini del modelo confirman el mismo perfil: mes_creacion, origen, dia_creacion "
        "y vehiculo_interes concentran la mayor parte del peso."
    )

    doc.add_heading("8.2 Patrones clave detectados por SHAP", level=2)
    patterns = [
        "mes_creacion bajo (enero-febrero) empuja hacia Hot; meses más altos empujan hacia Cold.",
        "Leads de madrugada/noche tienen SHAP positivo: mayor intención de compra genuina.",
        "Ciertos vehículos y orígenes concentran mayor probabilidad Hot de forma consistente.",
        "El fin de semana (es_fin_de_semana = 1) tiene un leve efecto positivo hacia Hot.",
    ]
    for p in patterns:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("8.3 Interpretación recomendada", level=2)
    doc.add_paragraph(
        "SHAP debe usarse para explicar por qué un lead sube o baja en probabilidad combinando "
        "contexto temporal, vehículo, origen, concesionario y formulario. La salida de la app permite "
        "explicar una observación individual sin usar terminología estadística avanzada."
    )

    # ================================================================
    # 9. LIMITACIONES, ÉTICA Y SEGURIDAD
    # ================================================================
    doc.add_heading("9. Limitaciones, ética y seguridad", level=1)

    doc.add_heading("9.1 Limitaciones vigentes", level=2)
    limitations = [
        "El umbral operativo correcto es 0.325; toda referencia a 0.35 se considera legacy.",
        "La distribución y calidad del lead pueden cambiar si el proceso comercial cambia; el modelo debe revalidarse ante nuevos cortes o nuevas reglas de negocio.",
        "Concesionarios o categorías nuevas se manejan con fallback: categorías no vistas se agrupan como 'otros' y concesionarios no vistos usan media global.",
        "Pendiente metodológico: evaluar si conservar simultáneamente nombre_formulario y campaña aporta información neta o solo redundancia.",
        "La app no usa PII como features. Esto reduce riesgo de privacidad, pero también limita señales potenciales provenientes de texto libre.",
        "No existe despliegue cloud activo; la app corre local.",
    ]
    for l in limitations:
        doc.add_paragraph(l, style="List Bullet")

    doc.add_heading("9.2 Controles de ética y seguridad", level=2)
    controls = [
        "No se usan nombre, teléfono, correo ni otros datos personales para predecir.",
        "Decisiones de leakage cerradas: exclusión de chatbot, exclusión de Arkana/Oroch preclasificados y exclusión de meses anómalos abril-junio.",
        "La persistencia local se limita a SQLite de historial; no existe despliegue cloud activo.",
        "Cada predicción puede auditarse con SHAP, reduciendo el carácter de caja negra.",
    ]
    for c in controls:
        doc.add_paragraph(c, style="List Bullet")

    # ================================================================
    # 10. IMPLEMENTACIÓN TÉCNICA
    # ================================================================
    doc.add_heading("10. Implementación técnica", level=1)
    doc.add_paragraph(
        "Detalle de las tecnologías y componentes usados en cada capa del sistema."
    )
    add_styled_table(doc,
        ["Capa", "Tecnología", "Detalle"],
        [
            ["Backend / lógica", "Python 3.13",
             "Pipeline de inferencia en streamlit_app/core/inference.py. "
             "Entrenamiento en notebooks con scikit-learn."],
            ["Frontend / UI", "Streamlit + Plotly",
             "App interactiva con 3 páginas, gráficas Plotly, "
             "CSS custom para temas claro/oscuro."],
            ["Modelo", "GradientBoostingClassifier (.joblib)",
             "Archivo models/best_model.joblib serializado con joblib. "
             "Artefactos de preprocesamiento en models/preprocessing_config.joblib."],
            ["Base de datos", "SQLite (local)",
             "Persistencia de historial de sesiones en data/memory/history.db "
             "mediante streamlit_app/core/memory.py."],
            ["Explicabilidad", "SHAP (TreeExplainer)",
             "Explicaciones globales y por lead individual en "
             "streamlit_app/core/shap_explainer.py."],
            ["Despliegue", "Local (sin cloud)",
             "Ejecución local con Streamlit. No existe despliegue cloud activo. "
             "Comando: .\\.venv\\Scripts\\python.exe -m streamlit run streamlit_app/app.py"],
        ],
    )

    # ================================================================
    # 11. RESULTADOS Y VALOR
    # ================================================================
    doc.add_heading("11. Resultados y valor", level=1)
    doc.add_paragraph(
        "El modelo reduce significativamente el tiempo de evaluación manual de leads al clasificar "
        "automáticamente cada prospecto como Hot o Cold con una probabilidad calibrada. Esto permite "
        "al equipo comercial de Renault México priorizar su atención en los leads con mayor "
        "probabilidad de conversión."
    )
    value_points = [
        "Reducción del tiempo de evaluación: el equipo comercial deja de revisar manualmente "
        "cada lead y se enfoca en los clasificados como Hot, reduciendo el esfuerzo operativo.",
        "Mejora en la tasa de contacto efectivo: al priorizar leads Hot (Recall 69.9%), "
        "se contactan más prospectos con intención real de compra en menos tiempo.",
        "Transparencia en la decisión: cada predicción incluye una explicación SHAP que "
        "detalla qué variables influyeron y en qué dirección, facilitando la confianza del "
        "equipo comercial en el modelo.",
        "Auditoría y trazabilidad: el historial de sesiones en SQLite permite revisar "
        "predicciones pasadas, comparar tendencias de % Hot entre lotes y exportar resultados.",
        "Reproducibilidad completa: el pipeline desde datos crudos hasta predicción está "
        "documentado en 5 notebooks secuenciales (00→04) y puede re-ejecutarse íntegramente.",
    ]
    for v in value_points:
        doc.add_paragraph(v, style="List Bullet")

    # ================================================================
    # 12. MANUAL DE USO
    # ================================================================
    doc.add_heading("12. Manual de uso", level=1)

    doc.add_heading("12.1 Requisitos de ejecución", level=2)
    doc.add_paragraph("Entorno virtual del repositorio con dependencias instaladas (ver requirements.txt).")
    doc.add_paragraph("Comando de arranque: .\\.venv\\Scripts\\python.exe -m streamlit run streamlit_app/app.py")

    doc.add_heading("12.2 Flujo de uso", level=2)
    flow_steps = [
        "Abrir la página 'Predicción de Leads'.",
        "Subir un Excel exportado desde Salesforce. El parser identifica columnas por keywords semánticos; no exige 28 columnas exactas siempre que estén presentes las columnas necesarias (fecha, formulario, campaña, origen de creación, vehículo, concesionario, origen).",
        "Revisar el banner de resumen (X leads procesados: Y Hot, Z Cold) y la grilla de resultados con filtros interactivos.",
        "Abrir el detalle individual de un lead para ver la explicación SHAP (qué features empujaron la predicción).",
        "Ir a 'Reporte del Modelo' para revisar los 5 tabs: Data Engineering, EDA, Feature Engineering, Modelado y Evaluación. Cada gráfica incluye explicación de cómo leerla, hallazgos y diferencias V1 vs V2.",
        "Ir a 'Historial' para reabrir sesiones guardadas en SQLite, ver tendencia de % Hot y exportar resultados previos.",
    ]
    for i, s in enumerate(flow_steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    doc.add_heading("12.3 Salida principal", level=2)
    outputs = [
        "CSV de predicciones descargable.",
        "Visualización de probabilidad Hot/Cold por lead.",
        "Explicación SHAP individual por lead.",
        "Sesión guardada en data/memory/history.db para auditoría.",
    ]
    for o in outputs:
        doc.add_paragraph(o, style="List Bullet")

    # ================================================================
    # 11. ESTRUCTURA DEL REPOSITORIO
    # ================================================================
    doc.add_heading("13. Estructura del repositorio", level=1)
    tree_text = """predictivo_ml_leads/
├── data/
│   ├── raw/                           → Insumos crudos y snapshots locales
│   ├── processed/                     → leads_cleaned.csv, X_train_v2.csv, X_test_v2.csv, y_train_v2.csv, y_test_v2.csv
│   ├── memory/history.db              → Historial SQLite de la app
│   └── test/                          → Archivos de prueba para inferencia
├── models/
│   ├── best_model.joblib              → GradientBoostingClassifier desplegado
│   └── preprocessing_config.joblib    → Artefactos de preprocesamiento (umbral legacy 0.35)
├── notebooks/
│   ├── 00_data_engineering_v2_2025.ipynb
│   ├── 01_exploratory_data_analysis_v2.ipynb
│   ├── 02_feature_engineering_v2.ipynb
│   ├── 03_modeling_v2.ipynb
│   ├── 04_evaluation_v2.ipynb
│   └── v1/                            → Notebooks legacy de referencia histórica
├── streamlit_app/
│   ├── app.py                         → Entrada principal de la app
│   ├── core/
│   │   ├── inference.py               → Pipeline de inferencia vigente
│   │   ├── memory.py                  → Persistencia SQLite
│   │   ├── shap_explainer.py          → Explicación SHAP por lead
│   │   └── theme.py                   → CSS custom para temas claro/oscuro
│   └── components/
│       ├── upload.py                  → Widget de carga de archivos
│       ├── results_grid.py            → Grilla de resultados con filtros
│       ├── lead_detail.py             → Detalle individual con SHAP
│       ├── model_report.py            → Reporte completo del modelo (5 tabs)
│       └── history.py                 → Historial de sesiones
├── src/
│   ├── predict.py                     → Flujo CLI / soporte programático
│   ├── save_artifacts.py              → Generación de preprocessing_config.joblib
│   └── train.py                       → Script de entrenamiento (legacy RF)
├── graphify-out/                      → Análisis de grafos del código
├── PENDIENTES_VALIDACION.md           → Estado de validaciones de leakage
├── CONTEXTO.md                        → Contexto general del proyecto
├── requirements.txt                   → Dependencias Python
└── README.md"""

    p = doc.add_paragraph()
    run = p.add_run(tree_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)

    # ================================================================
    # 12. CHECKLIST DE ENTREGA
    # ================================================================
    doc.add_heading("14. Checklist de entrega", level=1)
    add_styled_table(doc,
        ["Elemento", "Estado / Detalle"],
        [
            ["Documento técnico", "Actualizado (v2.2) con narrativa V2 coherente con el repo y el modelo desplegado."],
            ["Repositorio", "https://github.com/JocDLC/predictivo_ml_leads"],
            ["Video explicativo", "Pendiente de insertar URL final."],
            ["Demo", "App local con Streamlit; sin despliegue cloud activo."],
            ["Notebooks V2", "5 notebooks (00-04) re-ejecutados y verificados (Mayo 2026)."],
            ["Modelo desplegado", "GradientBoostingClassifier en models/best_model.joblib."],
        ],
    )

    doc.add_heading("Notas del estudiante", level=2)
    notes = [
        "El modelo vigente documentado y desplegado es GradientBoostingClassifier; se elimina la narrativa anterior de Random Forest.",
        "El umbral operativo vigente es 0.325.",
        "Métricas actuales en test a 0.325: ROC-AUC 0.7828, Accuracy 0.6947, Precision 0.5604, Recall 0.6992 y F1 0.6221.",
        "Validaciones de leakage cerradas: exclusión de MX_LEAD_QUALIF, exclusión de Arkana/Oroch preclasificados y exclusión de chatbot.",
        "Pendiente abierto: evaluar redundancia nombre_formulario vs campaña para un futuro reentrenamiento; el modelo hoy desplegado aún usa ambas.",
        "El pipeline reproducible corresponde a la secuencia 00→04 de notebooks V2 y a la app Streamlit que consume best_model.joblib.",
        "Todos los notebooks V2 fueron re-ejecutados y verificados en Mayo 2026.",
    ]
    for n in notes:
        doc.add_paragraph(n, style="List Bullet")

    # --- Save ---
    doc.save(OUTPUT_PATH)
    print(f"Documento generado: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
